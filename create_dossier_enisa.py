"""
Script to generate the ENISA Dossier as a professional Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color Palette ──
NAVY = RGBColor(0x0B, 0x1D, 0x3A)       # Dark navy
GOLD = RGBColor(0xC8, 0x96, 0x2E)        # Professional gold
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "0B1D3A"
TABLE_ALT_BG = "F7F7F7"
GOLD_HEX = "C8962E"
ACCENT_BLUE = RGBColor(0x1A, 0x56, 0xDB)

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK_GRAY

# ── Style Definitions ──
# Heading 1 - Section titles
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(22)
h1_style.font.color.rgb = NAVY
h1_style.font.bold = True
h1_style.paragraph_format.space_before = Pt(36)
h1_style.paragraph_format.space_after = Pt(12)
h1_style.paragraph_format.keep_with_next = True

# Heading 2 - Subsection titles
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(15)
h2_style.font.color.rgb = GOLD
h2_style.font.bold = True
h2_style.paragraph_format.space_before = Pt(24)
h2_style.paragraph_format.space_after = Pt(8)

# Heading 3
h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(12)
h3_style.font.color.rgb = NAVY
h3_style.font.bold = True
h3_style.paragraph_format.space_before = Pt(16)
h3_style.paragraph_format.space_after = Pt(6)


# ── Helper Functions ──

def add_gold_line():
    """Add a gold horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Create a border line
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{GOLD_HEX}"/></w:pBdr>')
    pPr.append(pBdr)


def add_navy_line():
    """Add a navy horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{TABLE_HEADER_BG}"/></w:pBdr>')
    pPr.append(pBdr)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(headers, rows, col_widths=None):
    """Create a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Style borders to be subtle
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.bold = True
        # Cell padding
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="60" w:type="dxa"/><w:bottom w:w="60" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
            # Bold first column
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = NAVY
            # Cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="50" w:type="dxa"/><w:bottom w:w="50" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>')
            tcPr.append(tcMar)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table


def add_body(text, bold_prefix=None, italic=False, space_after=Pt(6)):
    """Add body text with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(15)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    if italic:
        run.font.italic = True
    return p


def add_bullet(text, bold_prefix=None, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.clear()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p


def add_callout(text, icon=""):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="{GOLD_HEX}"/></w:pBdr>')
    pPr.append(pBdr)
    # Shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FDF8ED" w:val="clear"/>')
    pPr.append(shading)
    if icon:
        run = p.add_run(f"{icon} ")
        run.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.font.italic = True
    return p


def add_pending_box(text):
    """Add a red-bordered pending info box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="CC3333"/></w:pBdr>')
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF0F0" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run("PENDIENTE: ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    return p


def add_section_number(number):
    """Add a large section number indicator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(number)
    run.font.name = 'Calibri'
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(0xE8, 0xE0, 0xD0)
    run.font.bold = True
    return p


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

# Top spacing
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TUMAI")
run.font.name = 'Calibri'
run.font.size = Pt(52)
run.font.color.rgb = NAVY
run.font.bold = True

add_gold_line()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Plataforma de Automatizacion con IA para Empresas")
run.font.name = 'Calibri'
run.font.size = Pt(18)
run.font.color.rgb = MEDIUM_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Automation as a Product")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = GOLD
run.font.italic = True

add_gold_line()

# Spacing
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Document info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DOSSIER PARA SOLICITUD ENISA")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prestamo Participativo via Unidad de Grandes Empresas (UGE)")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = MEDIUM_GRAY

for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Marzo 2026")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = MEDIUM_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento confidencial")
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.font.italic = True

# Page break
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("INDICE")
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.color.rgb = NAVY
run.font.bold = True

add_navy_line()

toc_items = [
    ("01", "Resumen Ejecutivo"),
    ("02", "Descripcion Detallada del Proyecto"),
    ("03", "Analisis de Mercado"),
    ("04", "Modelo de Negocio"),
    ("05", "Plan Financiero (3 anos)"),
    ("06", "Impacto Economico en Espana"),
    ("07", "Perfil del Emprendedor"),
    ("", "Anexos Recomendados"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(2.0))
    if num:
        run = p.add_run(num)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.color.rgb = GOLD
        run.font.bold = True
        run = p.add_run("\t")
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY if num else MEDIUM_GRAY

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 1: RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════════

add_section_number("01")
doc.add_heading("RESUMEN EJECUTIVO", level=1)
add_navy_line()

doc.add_heading("El problema: la IA es inaccesible para las PyMEs", level=2)
add_body(
    "Las pequenas y medianas empresas (PyMEs) del sector inmobiliario y servicios enfrentan un problema critico: "
    "la automatizacion con inteligencia artificial es inaccesible para ellas. Las herramientas existentes (N8N, Make, Zapier) "
    "son horizontales, requieren configuracion tecnica avanzada, no incluyen capacidades nativas de IA (vision artificial, "
    "procesamiento de voz, agentes inteligentes) y obligan al usuario a construir todo desde cero."
)
add_body(
    "El resultado: solo las grandes empresas con equipos tecnicos pueden beneficiarse de la IA, mientras que las PyMEs "
    "quedan fuera de la revolucion tecnologica mas importante de la decada. En Espana, donde el 99,8% del tejido empresarial "
    "son PyMEs, esta brecha tecnologica amenaza la competitividad del ecosistema."
)

doc.add_heading("La solucion: Tumai", level=2)
add_body(
    "Tumai es una plataforma de \"Automatizacion como Producto\" (Automation as a Product) — una biblioteca de "
    "35 modulos de automatizacion pre-construidos, probados en produccion y potenciados por IA, que pueden componerse "
    "como bloques LEGO para crear soluciones empresariales personalizadas en dias, no meses.",
    bold_prefix=""
)

add_body("A diferencia de las herramientas existentes, Tumai integra de forma nativa:")
add_bullet("Vision artificial (GPT-4 Vision) para inspeccion de activos y clasificacion de imagenes", bold_prefix="Vision AI: ")
add_bullet("Procesamiento de voz (Whisper) para entrada de datos manos libres en campo", bold_prefix="Voz: ")
add_bullet("Agentes inteligentes con razonamiento (LLM function calling) para estimacion de costes, analisis de precios y asistencia de datos", bold_prefix="Agentes IA: ")
add_bullet("Modulos verticales pre-configurados por industria (inmobiliario, financiero, servicios, contabilidad)", bold_prefix="Verticales: ")

add_callout("El cliente no programa: solo configura variables. Onboarding en 1-3 dias frente a semanas de desarrollo custom.")

doc.add_heading("Por que ahora", level=2)
add_bullet("Madurez de la IA generativa: GPT-4o, Whisper y los modelos de vision han alcanzado calidad de produccion en 2024-2025", bold_prefix="1. ")
add_bullet("Regulacion europea favorable: El AI Act europeo incentiva el uso responsable de IA en empresas", bold_prefix="2. ")
add_bullet("Brecha digital en PyMEs: La digitalizacion post-COVID ha acelerado la urgencia", bold_prefix="3. ")
add_bullet("Ventana de oportunidad: No existe ninguna plataforma que combine automatizacion vertical + IA nativa + modulos pre-construidos", bold_prefix="4. ")

doc.add_heading("Ventaja competitiva", level=2)

add_styled_table(
    headers=["Aspecto", "Competidores (N8N/Make/Zapier)", "Tumai"],
    rows=[
        ["Agentes IA", "Solo plugins", "Nativos en el core"],
        ["Vision artificial", "No", "Si (GPT-4 Vision)"],
        ["Procesamiento de voz", "No", "Si (Whisper)"],
        ["LLM Function Calling", "Limitado", "Funcion central"],
        ["Configuracion", "DIY (hazlo tu mismo)", "Listo en minutos"],
        ["Expertise vertical", "Solo horizontal", "Especifico por industria"],
        ["Probado en produccion", "Marketplace", "Casos reales con clientes"],
    ]
)

add_callout(
    "Ventaja clave: Los 35 modulos ya estan probados en produccion con un cliente real (Maninos Homes LLC, Texas), "
    "procesando operaciones financieras, contratos, pagos y agentes de IA con +672 commits de desarrollo iterativo."
)

doc.add_heading("Impacto en Espana", level=2)
add_bullet("Plan de contratacion de 8-12 personas en los primeros 3 anos", bold_prefix="Creacion de empleo: ")
add_bullet("Tecnologia de IA validada internacionalmente, ahora aplicada al tejido empresarial espanol", bold_prefix="Transferencia tecnologica: ")
add_bullet("Hacer accesible la IA a miles de PyMEs espanolas", bold_prefix="Democratizacion de la IA: ")
add_bullet("Plataforma disenada para escalar a Latinoamerica y Europa desde Espana", bold_prefix="Exportacion digital: ")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 2: DESCRIPCION DETALLADA DEL PROYECTO
# ════════════════════════════════════════════════════════════════

add_section_number("02")
doc.add_heading("DESCRIPCION DETALLADA DEL PROYECTO", level=1)
add_navy_line()

doc.add_heading("Producto y servicio", level=2)
add_body(
    "Tumai es una plataforma SaaS de automatizacion empresarial potenciada por IA, organizada en modulos "
    "independientes y componibles. El producto se estructura en tres capas:"
)

doc.add_heading("Capa 1: Biblioteca de Automatizaciones (35 modulos)", level=3)

add_styled_table(
    headers=["Categoria", "Modulos", "Innovacion clave"],
    rows=[
        ["Pagos y Cobros", "8", "Maquina de estados para ciclo de pagos, scoring de morosidad, analisis de inversiones"],
        ["Contabilidad", "5", "Clasificador de transacciones con GPT-4o, reconciliacion automatica, motor de KPIs"],
        ["Agentes IA", "7", "Vision (GPT-4), Voz (Whisper), Function calling, estimacion costes/precios con LLM"],
        ["Comunicacion", "3", "Cola de emails, recordatorios de pago, alertas de morosidad"],
        ["Data Collection", "3", "Scraper de APIs, automatizacion de navegador (Playwright), extraccion con vision"],
        ["Documentos", "4", "Generacion de PDFs, almacenamiento automatico, contratos, reportes periodicos"],
        ["Analytics y Reglas", "3", "Prediccion de precios, motor de reglas de negocio, analisis financiero"],
        ["Infraestructura", "2", "Scheduler centralizado, sincronizacion entre sistemas"],
    ]
)

doc.add_heading("Capa 2: Plataforma Visual de Composicion", level=3)
add_bullet("Editor de nodos tipo canvas basado en React Flow v12")
add_bullet("Drag-and-drop para disenar flujos de automatizacion")
add_bullet("Visualizacion en tiempo real de dependencias entre modulos")
add_bullet("Panel de detalle con variables configurables, dependencias y casos de uso")

doc.add_heading("Capa 3: Paquetes Verticales por Industria", level=3)

add_styled_table(
    headers=["Vertical", "Automatizaciones", "Cliente objetivo"],
    rows=[
        ["Alquiler de propiedades", "20 modulos", "Empresas de alquiler"],
        ["Compra-Reforma-Venta", "35 modulos (todos)", "Flippers inmobiliarios"],
        ["Negocio de servicios", "15 modulos", "HVAC, limpieza, mantenimiento"],
        ["Financiacion / Prestamos", "18 modulos", "Empresas de prestamos, microfinanzas"],
        ["Contabilidad", "12 modulos", "Despachos contables, PyMEs"],
    ]
)

doc.add_heading("Tecnologia utilizada", level=2)

add_styled_table(
    headers=["Capa", "Tecnologia", "Detalles"],
    rows=[
        ["Backend", "Python 3.12 + FastAPI", "Framework asincrono de alto rendimiento"],
        ["Frontend", "Next.js 14 + React + TypeScript", "React Flow v12 para editor visual de nodos"],
        ["Base de datos", "PostgreSQL (Supabase)", "74+ migraciones, esquema robusto"],
        ["IA / ML", "OpenAI GPT-4o, GPT-4 Vision, Whisper", "Razonamiento, vision artificial, voz"],
        ["Orquestacion IA", "LangChain / LangGraph", "Agentes con herramientas y memoria"],
        ["Infraestructura", "Docker, Railway, Vercel", "Despliegue continuo automatizado"],
        ["Integraciones", "Stripe, Resend, Playwright", "Pagos, email, scraping web"],
        ["Observabilidad", "Logfire + Langfuse", "Monitorizacion de IA y aplicacion"],
    ]
)

doc.add_heading("Estado de desarrollo", level=2)
add_body("Estado actual: MVP funcional en produccion.", bold_prefix="")

add_styled_table(
    headers=["Indicador", "Valor"],
    rows=[
        ["Automatizaciones operativas", "35 en produccion con cliente real"],
        ["Commits de desarrollo", "+672 iteraciones"],
        ["Migraciones de base de datos", "74+ aplicadas"],
        ["Endpoints API activos", "30+"],
        ["Agentes de IA funcionales", "6 (costes, precios, fotos, voz, datos, inspeccion)"],
        ["Portales web desplegados", "3 en uso diario"],
    ]
)

doc.add_heading("Grado de innovacion", level=2)

doc.add_heading("Innovacion tecnologica", level=3)
add_bullet("Agentes de IA nativos: no son plugins anadidos; la IA esta integrada en el core de cada modulo")
add_bullet("Vision artificial aplicada: GPT-4 Vision inspecciona activos, clasifica fotos y extrae datos de capturas")
add_bullet("Procesamiento de voz en campo: Whisper permite entrada de datos manos libres durante inspecciones")

doc.add_heading("Innovacion de modelo de negocio", level=3)
add_bullet("\"Automation as a Product\": automatizaciones pre-construidas y probadas, no DIY")
add_bullet("Verticalizacion: paquetes especificos por industria con configuraciones optimizadas")

doc.add_heading("Innovacion de proceso", level=3)
add_bullet("Onboarding en 1-3 dias frente a semanas de desarrollo custom")
add_bullet("Modulos componibles: arquitectura LEGO que escala de 5 a 35 modulos")
add_bullet("Production-first: cada modulo nace de un caso real, no de una especificacion teorica")

add_callout(
    "No es una actividad profesional tradicional. Una consultoria de automatizacion cobra por proyecto y entrega "
    "codigo custom que no escala. Tumai es una plataforma de producto tecnologico SaaS: modulos estandarizados, "
    "configuracion sin codigo, escalables a miles de clientes simultaneamente."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 3: ANALISIS DE MERCADO
# ════════════════════════════════════════════════════════════════

add_section_number("03")
doc.add_heading("ANALISIS DE MERCADO", level=1)
add_navy_line()

doc.add_heading("Tamano del mercado", level=2)

doc.add_heading("TAM — Mercado global de automatizacion empresarial con IA", level=3)
add_body("$15.800 millones (2024), proyectado a $46.400 millones para 2029 (CAGR 24%). Fuentes: Gartner, McKinsey, Grand View Research.")

doc.add_heading("SAM — Automatizacion para PyMEs en Europa y Latinoamerica", level=3)
add_body("~$3.200 millones — PyMEs en sectores inmobiliario, financiero y servicios en Espana, Europa occidental y Latinoamerica. Espana cuenta con 3,2 millones de PyMEs con penetracion de automatizacion inferior al 15%.")

doc.add_heading("SOM — Primeros 3 anos", level=3)

add_styled_table(
    headers=["Periodo", "Clientes objetivo", "ARR estimado"],
    rows=[
        ["Ano 1", "20-30 clientes", "~300.000 EUR"],
        ["Ano 2", "80-120 clientes", "~1.200.000 EUR"],
        ["Ano 3", "250-400 clientes", "~3.500.000 EUR"],
    ]
)

doc.add_heading("Tendencias del sector", level=2)
add_bullet("Explosion de la IA generativa en empresas: 72% de empresas han adoptado alguna forma de IA en 2024 (McKinsey)", bold_prefix="1. ")
add_bullet("Automatizacion vertical sobre horizontal: las empresas prefieren soluciones especificas para su industria", bold_prefix="2. ")
add_bullet("Low-code/no-code en auge: el mercado crece al 25% anual", bold_prefix="3. ")
add_bullet("Regulacion favorable: el AI Act europeo incentiva la adopcion responsable", bold_prefix="4. ")
add_bullet("Digitalizacion de PyMEs: el 67% de PyMEs europeas han acelerado sus planes post-COVID", bold_prefix="5. ")

doc.add_heading("Competencia", level=2)

add_styled_table(
    headers=["Competidor", "Tipo", "Fortaleza", "Debilidad vs Tumai"],
    rows=[
        ["Zapier", "Horizontal, no-code", "Gran ecosistema de integraciones", "Sin IA nativa, sin verticalizacion, DIY"],
        ["Make (Integromat)", "Horizontal, visual", "Buen editor visual", "Sin agentes IA, sin modulos pre-construidos"],
        ["N8N", "Open source, horizontal", "Flexibilidad, self-hosted", "Requiere conocimiento tecnico, sin IA nativa"],
        ["UiPath / Automation Anywhere", "RPA enterprise", "Mercado enterprise consolidado", "Demasiado caro y complejo para PyMEs"],
        ["Consultorias IA", "Servicios custom", "Soluciones a medida", "No escalable, alto coste, dependencia"],
    ]
)

add_callout(
    "Espacio vacio que ocupa Tumai: no existe ninguna plataforma que combine automatizacion vertical + IA nativa "
    "(vision, voz, agentes) + modulos pre-construidos + precio accesible para PyMEs."
)

doc.add_heading("Ventaja competitiva clara", level=2)
add_bullet("35 modulos con +672 commits de iteracion real, no prototipos", bold_prefix="Probado en produccion: ")
add_bullet("Vision, voz y agentes inteligentes desde el diseno", bold_prefix="IA nativa: ")
add_bullet("1-3 dias de onboarding vs semanas de desarrollo", bold_prefix="Time-to-value: ")
add_bullet("Paquetes especificos por industria con best practices incluidas", bold_prefix="Verticalizacion: ")
add_bullet("Conocimiento profundo de procesos inmobiliarios y financieros codificado en los modulos", bold_prefix="Barrera de entrada: ")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 4: MODELO DE NEGOCIO
# ════════════════════════════════════════════════════════════════

add_section_number("04")
doc.add_heading("MODELO DE NEGOCIO", level=1)
add_navy_line()

doc.add_heading("Como genera ingresos Tumai", level=2)
add_body("Modelo SaaS (Software as a Service) con suscripcion mensual + fee de setup:")

doc.add_heading("Estructura de precios", level=3)

add_styled_table(
    headers=["Tier", "Automatizaciones incluidas", "Precio mensual"],
    rows=[
        ["Starter (Alquiler)", "20 modulos", "499 EUR/mes"],
        ["Pro (Alquiler + IA)", "20 modulos + agentes IA", "999 EUR/mes"],
        ["Enterprise (Buy-Renovate-Sell)", "35 modulos (todos)", "1.499 EUR/mes"],
        ["Servicios", "15 modulos", "499 EUR/mes"],
        ["Financiacion", "18 modulos", "799 EUR/mes"],
        ["Contabilidad", "12 modulos", "399 EUR/mes"],
        ["Custom Vertical", "Setup nuevo vertical", "2.999 EUR/mes + desarrollo"],
    ]
)

doc.add_heading("Ingresos adicionales", level=3)
add_bullet("Setup fee: 1.000-3.000 EUR por cliente (configuracion + formacion)")
add_bullet("Custom modules: Desarrollo de automatizaciones a medida (bajo demanda)")
add_bullet("Soporte premium: SLA garantizado para clientes enterprise")

doc.add_heading("Estructura de costes", level=2)

add_styled_table(
    headers=["Concepto", "Estimacion mensual (Ano 1)"],
    rows=[
        ["Infraestructura cloud (Railway, Vercel, Supabase)", "500-1.500 EUR"],
        ["APIs de IA (OpenAI GPT-4o, Whisper)", "500-2.000 EUR"],
        ["Salarios equipo (2-3 personas iniciales)", "8.000-15.000 EUR"],
        ["Marketing y ventas", "1.000-3.000 EUR"],
        ["Herramientas y licencias", "200-500 EUR"],
        ["Total estimado", "10.200-22.000 EUR/mes"],
    ]
)

doc.add_heading("Estrategia de pricing", level=2)
add_bullet("El ahorro que genera la automatizacion justifica ampliamente el coste mensual", bold_prefix="Basado en valor: ")
add_bullet("Entrar con un vertical basico (399-499 EUR) y expandir a modulos adicionales", bold_prefix="Land & expand: ")
add_bullet("15-20% de descuento por compromiso anual", bold_prefix="Descuentos anuales: ")
add_bullet("14 dias de prueba con modulos limitados", bold_prefix="Free trial: ")

doc.add_heading("Canales de captacion", level=2)
add_bullet("Enfoque en inmobiliarias, gestorias y empresas de servicios en Espana", bold_prefix="Venta directa: ")
add_bullet("Blog tecnico, casos de estudio (Maninos como caso flagship), webinars", bold_prefix="Content marketing: ")
add_bullet("Acuerdos con asociaciones sectoriales (APCE, AEI, camaras de comercio)", bold_prefix="Partnerships: ")
add_bullet("Descuentos por referidos entre clientes", bold_prefix="Referral program: ")
add_bullet("Posicionamiento como experto en automatizacion con IA", bold_prefix="LinkedIn + comunidades: ")
add_bullet("Partners integradores/consultoras que implementen Tumai en sus clientes", bold_prefix="Canal indirecto: ")

doc.add_heading("Metricas clave proyectadas", level=2)

add_styled_table(
    headers=["Metrica", "Objetivo Ano 1", "Objetivo Ano 3"],
    rows=[
        ["CAC (Coste de adquisicion)", "1.500-2.500 EUR", "800-1.200 EUR"],
        ["LTV (Valor de vida del cliente)", "12.000-18.000 EUR", "24.000-36.000 EUR"],
        ["LTV / CAC ratio", "5x - 8x", "20x - 30x"],
        ["Margen bruto", "65-70%", "75-85%"],
        ["Churn mensual", "< 5%", "< 3%"],
        ["ARPU (Ingreso medio por usuario)", "700 EUR/mes", "900 EUR/mes"],
    ]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 5: PLAN FINANCIERO
# ════════════════════════════════════════════════════════════════

add_section_number("05")
doc.add_heading("PLAN FINANCIERO (3 ANOS)", level=1)
add_navy_line()

doc.add_heading("Proyeccion de ingresos", level=2)

add_styled_table(
    headers=["Concepto", "Ano 1", "Ano 2", "Ano 3"],
    rows=[
        ["Clientes nuevos", "25", "80", "200"],
        ["Clientes acumulados (con churn)", "22", "85", "250"],
        ["ARPU mensual", "700 EUR", "800 EUR", "900 EUR"],
        ["Ingresos recurrentes (ARR)", "185.000 EUR", "816.000 EUR", "2.700.000 EUR"],
        ["Setup fees", "50.000 EUR", "120.000 EUR", "200.000 EUR"],
        ["Custom development", "30.000 EUR", "80.000 EUR", "150.000 EUR"],
        ["INGRESOS TOTALES", "265.000 EUR", "1.016.000 EUR", "3.050.000 EUR"],
    ]
)

doc.add_heading("Gastos operativos", level=2)

add_styled_table(
    headers=["Concepto", "Ano 1", "Ano 2", "Ano 3"],
    rows=[
        ["Salarios (3 -> 6 -> 12 personas)", "120.000 EUR", "300.000 EUR", "720.000 EUR"],
        ["Infraestructura cloud", "12.000 EUR", "36.000 EUR", "96.000 EUR"],
        ["APIs de IA", "15.000 EUR", "48.000 EUR", "120.000 EUR"],
        ["Marketing y ventas", "24.000 EUR", "72.000 EUR", "180.000 EUR"],
        ["Oficina y operaciones", "12.000 EUR", "24.000 EUR", "48.000 EUR"],
        ["Legal y contabilidad", "6.000 EUR", "12.000 EUR", "18.000 EUR"],
        ["Herramientas y licencias", "3.000 EUR", "6.000 EUR", "12.000 EUR"],
        ["Otros / contingencia", "8.000 EUR", "15.000 EUR", "30.000 EUR"],
        ["GASTOS TOTALES", "200.000 EUR", "513.000 EUR", "1.224.000 EUR"],
    ]
)

doc.add_heading("EBITDA estimado", level=2)

add_styled_table(
    headers=["Concepto", "Ano 1", "Ano 2", "Ano 3"],
    rows=[
        ["Ingresos", "265.000 EUR", "1.016.000 EUR", "3.050.000 EUR"],
        ["Gastos", "200.000 EUR", "513.000 EUR", "1.224.000 EUR"],
        ["EBITDA", "65.000 EUR", "503.000 EUR", "1.826.000 EUR"],
        ["Margen EBITDA", "24,5%", "49,5%", "59,9%"],
    ]
)

doc.add_heading("Necesidades de financiacion", level=2)

add_styled_table(
    headers=["Concepto", "Importe"],
    rows=[
        ["Desarrollo de plataforma (Fases 2-4)", "80.000 EUR"],
        ["Contratacion equipo inicial", "60.000 EUR"],
        ["Marketing de lanzamiento", "20.000 EUR"],
        ["Capital circulante (6 meses)", "60.000 EUR"],
        ["Infraestructura y herramientas", "10.000 EUR"],
        ["Legal (constitucion, PI, contratos)", "10.000 EUR"],
        ["Contingencia", "10.000 EUR"],
        ["TOTAL SOLICITADO", "250.000 EUR"],
    ]
)

add_pending_box(
    "Ajustar la cantidad solicitada segun los limites del programa ENISA aplicable: "
    "Jovenes Emprendedores (hasta 75K EUR) / Emprendedores (hasta 300K EUR) / Crecimiento (hasta 1,5M EUR). "
    "Confirmar con la UGE cual aplica."
)

doc.add_heading("Punto de equilibrio", level=2)
add_bullet("~17.000 EUR/mes de gastos fijos", bold_prefix="Breakeven mensual estimado: ")
add_bullet("~25 clientes a 700 EUR ARPU", bold_prefix="Clientes necesarios: ")
add_bullet("Mes 10-12 del Ano 1", bold_prefix="Punto de equilibrio estimado: ")

doc.add_heading("Hipotesis clave", level=2)
add_bullet("Churn mensual: 3-5% (conservador para SaaS B2B vertical)", bold_prefix="1. ")
add_bullet("Ciclo de venta: 30-60 dias para PyMEs, 60-90 dias para enterprise", bold_prefix="2. ")
add_bullet("Crecimiento organico: 40% de nuevos clientes por referral a partir del Ano 2", bold_prefix="3. ")
add_bullet("Expansion de ARPU: clientes anaden modulos adicionales (+15% ARPU anual)", bold_prefix="4. ")
add_bullet("Coste de APIs de IA: reduccion progresiva por competencia entre proveedores", bold_prefix="5. ")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 6: IMPACTO ECONOMICO EN ESPANA
# ════════════════════════════════════════════════════════════════

add_section_number("06")
doc.add_heading("IMPACTO ECONOMICO EN ESPANA", level=1)
add_navy_line()

doc.add_heading("Creacion de empleo prevista", level=2)

add_styled_table(
    headers=["Periodo", "Puestos nuevos", "Perfiles"],
    rows=[
        ["Ano 1", "3", "2 desarrolladores full-stack, 1 comercial/customer success"],
        ["Ano 2", "+3 adicionales", "1 desarrollador IA, 1 comercial, 1 soporte tecnico"],
        ["Ano 3", "+6 adicionales", "2 desarrolladores, 1 product manager, 1 marketing, 1 comercial, 1 operaciones"],
        ["TOTAL ANO 3", "12 empleados", "Equipo multidisciplinar con alto componente tecnologico"],
    ]
)

add_body("Tipo de empleo: cualificado, indefinido, con salarios competitivos. Perfiles STEM con especializacion en IA, desarrollo de software y producto.")

doc.add_heading("Inversion en Espana", level=2)
add_bullet("La empresa se constituye como S.L. espanola", bold_prefix="Sede fiscal y operativa: ")
add_bullet("250.000 EUR (financiacion ENISA) + fondos propios", bold_prefix="Inversion directa: ")
add_bullet("Salarios, oficina, servicios profesionales, proveedores locales", bold_prefix="Gasto recurrente: ")
add_bullet("Registrada en Espana, generando valor de activo intangible en el pais", bold_prefix="Propiedad intelectual: ")

doc.add_heading("Transferencia de conocimiento", level=2)
add_bullet(
    "La tecnologia ha sido desarrollada y validada con un cliente real en Texas (Maninos Homes LLC). "
    "35 automatizaciones probadas en produccion, 6 agentes de IA, 74+ migraciones — todo se transfiere al ecosistema espanol.",
    bold_prefix="De EE.UU. a Espana: "
)
add_bullet(
    "Conocimiento practico (no teorico) de como implementar GPT-4 Vision, Whisper y agentes inteligentes en procesos empresariales reales.",
    bold_prefix="IA aplicada: "
)
add_bullet(
    "El equipo contratado en Espana recibira formacion en tecnologias de IA de vanguardia.",
    bold_prefix="Formacion: "
)

doc.add_heading("Colaboracion con universidades y centros tecnologicos", level=2)
add_pending_box(
    "Identificar posibles colaboraciones con: UPM, UPC, UPV, UAM (programas de IA/ML), "
    "centros tecnologicos (CTIC, Barcelona Supercomputing Center, AI+), "
    "programas de practicas/becas, catedras de empresa o proyectos de I+D+i colaborativos."
)

doc.add_heading("Atraccion de inversion extranjera", level=2)
add_bullet("Maninos (Texas) como primer cliente que valida el modelo", bold_prefix="Cliente anchor internacional: ")
add_bullet("Espana como hub natural para Latinoamerica (idioma, zona horaria, conexiones)", bold_prefix="Mercado LATAM: ")
add_bullet("JME, K Fund, Samaipata, Nauta y ecosistema VC creciente", bold_prefix="Rondas futuras: ")

doc.add_heading("Efecto tractor", level=2)
add_bullet("Consultoras e integradores locales que implementen Tumai, generando empleo indirecto", bold_prefix="Ecosistema de partners: ")
add_bullet("Posibilidad de liberar componentes como open source", bold_prefix="Open source parcial: ")
add_bullet("Startup de IA con validacion internacional que escala desde Espana", bold_prefix="Caso replicable: ")
add_bullet("Cada cliente de Tumai es una PyME que se digitaliza, multiplicando el impacto", bold_prefix="Digitalizacion de PyMEs: ")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SECTION 7: PERFIL DEL EMPRENDEDOR
# ════════════════════════════════════════════════════════════════

add_section_number("07")
doc.add_heading("PERFIL DEL EMPRENDEDOR", level=1)
add_navy_line()

doc.add_heading("CV detallado", level=2)
add_pending_box(
    "Incluir: nombre completo, fecha de nacimiento, nacionalidad, "
    "formacion academica (universidad, master, certificaciones), "
    "experiencia laboral detallada (empresas, puestos, duracion, logros), LinkedIn URL."
)

doc.add_heading("Experiencia previa relevante", level=2)

doc.add_heading("Fundadora y CEO de Tumai", level=3)
add_bullet("Creacion de una plataforma de automatizacion con IA con 35 modulos probados en produccion")
add_bullet("Implementacion de 6 agentes de IA (vision, voz, razonamiento)")
add_bullet("Stack tecnologico completo: Python, FastAPI, Next.js, GPT-4, Whisper, LangChain")
add_bullet("Diseno de arquitectura multi-vertical para 5 industrias")

doc.add_heading("Proyecto Maninos Homes LLC (Cliente flagship, Texas, EE.UU.)", level=3)
add_bullet("Diseno y desarrollo de plataforma completa de gestion inmobiliaria (3 portales, 30+ APIs, 74+ tablas)")
add_bullet("Sistema de financiacion RTO (Rent-to-Own) con gestion de contratos, pagos e inversores")
add_bullet("Integracion de Stripe (pagos + KYC), Resend (email), Playwright (scraping)")
add_bullet("+672 commits de desarrollo iterativo en produccion")
add_bullet("Resultado: transformacion digital completa de la operacion (de Excel a plataforma con IA)")

doc.add_heading("Logros profesionales", level=2)
add_bullet("Plataforma SaaS con 35 automatizaciones en produccion")
add_bullet("6 agentes de IA funcionales en entorno empresarial real")
add_bullet("Sistema financiero completo (contratos, pagos, inversores, contabilidad)")
add_bullet("Validacion internacional del producto con cliente en EE.UU.")

doc.add_heading("Experiencia internacional", level=2)
add_bullet("Trabajo con clientes en Estados Unidos (Texas)")
add_bullet("Desarrollo de producto con enfoque multi-mercado (EE.UU., Espana, Latinoamerica)")
add_bullet("Plataforma disenada desde el inicio para ser multi-idioma y multi-region")

doc.add_heading("Logros adicionales", level=2)
add_pending_box(
    "Indicar si tienes: exits anteriores (venta de empresas), publicaciones (articulos, papers, conferencias), "
    "patentes, financiacion previa (inversores, subvenciones, premios), premios o reconocimientos, "
    "participacion en aceleradoras o incubadoras, mentoring o advisory, certificaciones tecnicas relevantes."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# ANNEXES
# ════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("ANEXOS RECOMENDADOS")
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.color.rgb = NAVY
run.font.bold = True

add_navy_line()

add_body("Se recomienda adjuntar al dossier:")

add_bullet("Capturas de pantalla de la plataforma Tumai (canvas visual, panel de automatizaciones)", bold_prefix="1. ")
add_bullet("Capturas de pantalla de Maninos en produccion (portales, dashboards, agentes IA)", bold_prefix="2. ")
add_bullet("Demo en video (2-3 minutos) mostrando el flujo de configuracion", bold_prefix="3. ")
add_bullet("Diagrama de arquitectura tecnica", bold_prefix="4. ")
add_bullet("Listado completo de las 35 automatizaciones con descripcion", bold_prefix="5. ")
add_bullet("Carta de referencia de Maninos Homes LLC", bold_prefix="6. ")
add_bullet("CV detallado del emprendedor", bold_prefix="7. ")
add_bullet("Plan de hitos con milestones trimestrales", bold_prefix="8. ")

# ── Footer line ──
doc.add_paragraph()
add_gold_line()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tumai  |  Automatizacion como Producto  |  Marzo 2026")
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento confidencial — Solicitud ENISA via UGE")
run.font.name = 'Calibri'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)


# ── Save ──
output_path = "/Users/mariasebares/Documents/RAMA_AI/maninos-ai/docs/DOSSIER_ENISA_TUMAI.docx"
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
